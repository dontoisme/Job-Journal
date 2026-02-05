"""Resume import - parse existing resumes and link to corpus entries.

This module:
1. Parses PDF/DOCX resume files to extract text
2. Identifies bullet points and sections
3. Matches bullets to corpus entries using fuzzy matching
4. Creates resume records with entry linkages
5. Flags unmatched bullets for review
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from jj.corpus import find_matching_entry, find_all_matching_entries, validate_bullet
from jj.db import (
    create_resume,
    create_resume_entry,
    create_resume_section,
    create_corpus_suggestion,
    get_resume_by_filepath,
    get_roles,
    find_role_by_company_title,
)


@dataclass
class ParsedResume:
    """Structured representation of a parsed resume."""

    filepath: Path
    filename: str
    variant: Optional[str] = None
    target_company: Optional[str] = None
    target_role: Optional[str] = None
    summary: Optional[str] = None
    skills_sections: list[dict[str, Any]] = field(default_factory=list)
    experience_sections: list[dict[str, Any]] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    raw_text: str = ""


@dataclass
class ImportResult:
    """Result of a resume import operation."""

    resume_id: Optional[int] = None
    filepath: str = ""
    entries_linked: int = 0
    entries_unmatched: int = 0
    unmatched_bullets: list[str] = field(default_factory=list)
    suggestions_created: int = 0
    skipped: bool = False
    error: Optional[str] = None


def extract_text_from_pdf(filepath: Path) -> str:
    """Extract text content from a PDF file."""
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except ImportError:
        # Fallback: try pdfplumber
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError(
                "PDF parsing requires pypdf or pdfplumber. "
                "Install with: pip install pypdf"
            )


def extract_text_from_docx(filepath: Path) -> str:
    """Extract text content from a DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError(
            "DOCX parsing requires python-docx. "
            "Install with: pip install python-docx"
        )


def extract_bullets_from_text(text: str) -> list[str]:
    """
    Extract bullet points from resume text.

    Looks for common bullet patterns:
    - Lines starting with bullet characters (•, -, *, ►, ●)
    - Lines starting with action verbs (common resume pattern)
    """
    bullets = []

    # Clean up text - PDF extraction often splits words across lines
    # Join lines that end mid-word
    text = re.sub(r'(\w)\n(\w)', r'\1\2', text)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)

    # Action verbs that typically start resume bullets (capitalized as they appear)
    action_verbs = [
        "Led", "Built", "Developed", "Created", "Designed", "Implemented",
        "Managed", "Owned", "Drove", "Increased", "Reduced", "Improved",
        "Delivered", "Launched", "Scaled", "Established", "Integrated",
        "Collaborated", "Partnered", "Defined", "Guided", "Applied",
        "Configured", "Leveraged", "Coordinated", "Consolidated",
        "Achieved", "Accelerated", "Generated", "Transformed",
        "Architected", "Spearheaded", "Optimized", "Streamlined",
        "Set", "Shipped", "Engaged", "Mentored", "Hands",  # Additional common starters
    ]
    action_pattern = "|".join(action_verbs)

    # Bullet character patterns - extract content starting with action verbs
    # Simple approach: find ● followed by capitalized word and capture until next ●
    bullet_pattern = re.compile(r'●\s*([A-Z][a-z]+\s+[^●]+)')
    matches = bullet_pattern.findall(text)

    # Filter to only keep those starting with action verbs
    for match in matches:
        first_word = match.split()[0] if match.split() else ""
        if first_word in action_verbs:
            bullets.append(match.strip())

    # Also look for lines that start with action verbs (resume bullets often do)
    action_verbs = [
        "Led", "Built", "Developed", "Created", "Designed", "Implemented",
        "Managed", "Owned", "Drove", "Increased", "Reduced", "Improved",
        "Delivered", "Launched", "Scaled", "Established", "Integrated",
        "Collaborated", "Partnered", "Defined", "Guided", "Applied",
        "Configured", "Leveraged", "Coordinated", "Consolidated",
    ]

    for line in text.split("\n"):
        line = line.strip()
        if not line or len(line) < 20:
            continue

        # Check if starts with action verb
        for verb in action_verbs:
            if line.startswith(verb + " ") or line.startswith(verb + ","):
                if line not in bullets:
                    bullets.append(line)
                break

    # Clean up bullets
    cleaned = []
    for bullet in bullets:
        bullet = bullet.strip()
        # Remove bullet character if still present
        bullet = re.sub(r'^[•\-\*►◦]\s*', '', bullet)
        # Skip short or empty bullets
        if len(bullet) >= 20:
            cleaned.append(bullet)

    # Deduplicate while preserving order
    seen = set()
    unique = []
    for b in cleaned:
        if b not in seen:
            seen.add(b)
            unique.append(b)

    return unique


def parse_filename_metadata(filename: str) -> dict[str, Optional[str]]:
    """
    Extract metadata from resume filename.

    Expected patterns:
    - "Jane Smith - Principal PM - Company - Resume.docx"
    - "Jane Smith - Growth PM - 2025.pdf"
    - "Resume - CompanyName.pdf"
    """
    metadata = {
        "target_company": None,
        "target_role": None,
        "variant": None,
    }

    # Try common pattern: Name - Role - Company - Resume
    parts = filename.replace(".pdf", "").replace(".docx", "").split(" - ")

    if len(parts) >= 3:
        # Assume: Name - Role - Company - ...
        role_part = parts[1].lower()
        company_part = parts[2] if len(parts) > 2 else None

        if company_part and company_part.lower() not in ["resume", "cv"]:
            metadata["target_company"] = company_part

        metadata["target_role"] = parts[1]

        # Detect variant from role
        if "growth" in role_part:
            metadata["variant"] = "growth"
        elif "ai" in role_part or "agentic" in role_part:
            metadata["variant"] = "ai-agentic"
        elif "health" in role_part:
            metadata["variant"] = "health-tech"
        elif "consumer" in role_part:
            metadata["variant"] = "consumer"

    return metadata


def parse_resume(filepath: Path) -> ParsedResume:
    """
    Parse a resume file and extract structured content.

    Args:
        filepath: Path to PDF or DOCX file

    Returns:
        ParsedResume with extracted content
    """
    # Extract raw text based on file type
    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        raw_text = extract_text_from_pdf(filepath)
    elif suffix in [".docx", ".doc"]:
        raw_text = extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    # Extract metadata from filename
    metadata = parse_filename_metadata(filepath.name)

    # Extract bullets
    bullets = extract_bullets_from_text(raw_text)

    return ParsedResume(
        filepath=filepath,
        filename=filepath.name,
        variant=metadata["variant"],
        target_company=metadata["target_company"],
        target_role=metadata["target_role"],
        bullets=bullets,
        raw_text=raw_text,
    )


def import_resume(
    filepath: Path,
    match_threshold: float = 0.85,
    create_suggestions: bool = True,
) -> ImportResult:
    """
    Import a resume file and link bullets to corpus entries.

    Args:
        filepath: Path to resume file
        match_threshold: Minimum fuzzy match score for linking
        create_suggestions: Whether to create suggestions for unmatched bullets

    Returns:
        ImportResult with import statistics
    """
    result = ImportResult(filepath=str(filepath))

    # Check if already imported
    existing = get_resume_by_filepath(str(filepath))
    if existing:
        result.skipped = True
        result.resume_id = existing["id"]
        return result

    # Parse the resume
    try:
        parsed = parse_resume(filepath)
    except Exception as e:
        result.error = str(e)
        return result

    # Create resume record
    resume_id = create_resume(
        filename=parsed.filename,
        filepath=str(filepath),
        variant=parsed.variant,
        target_company=parsed.target_company,
        target_role=parsed.target_role,
    )
    result.resume_id = resume_id

    # Get roles for matching
    roles = get_roles()
    roles_by_company = {}
    for role in roles:
        company_key = role["company"].lower()
        roles_by_company[company_key] = role

    # Match bullets to corpus entries
    position = 0
    for bullet in parsed.bullets:
        match = find_matching_entry(bullet, threshold=match_threshold)

        if match:
            # Get the role_id from the matched entry
            role_id = match.get("role_id")
            if role_id:
                create_resume_entry(
                    resume_id=resume_id,
                    entry_id=match["id"],
                    role_id=role_id,
                    position=position,
                )
                result.entries_linked += 1
                position += 1
        else:
            result.entries_unmatched += 1
            result.unmatched_bullets.append(bullet)

            # Create suggestion for unmatched bullet
            if create_suggestions:
                # Try to identify which role this might belong to
                suggested_role_id = None
                # Look for company names in the bullet
                for company_key, role in roles_by_company.items():
                    if company_key in bullet.lower():
                        suggested_role_id = role["id"]
                        break

                create_corpus_suggestion(
                    gap_type="unmatched_bullet",
                    theme="import",
                    suggestion=f"Unmatched bullet from imported resume: {bullet[:200]}...",
                    resume_id=resume_id,
                    suggested_role_id=suggested_role_id,
                )
                result.suggestions_created += 1

    return result


def import_directory(
    directory: Path,
    recursive: bool = True,
    match_threshold: float = 0.85,
) -> list[ImportResult]:
    """
    Import all resumes from a directory.

    Args:
        directory: Directory path to scan
        recursive: Whether to scan subdirectories
        match_threshold: Minimum fuzzy match score for linking

    Returns:
        List of ImportResult for each file
    """
    results = []

    pattern = "**/*" if recursive else "*"
    extensions = [".pdf", ".docx"]

    for ext in extensions:
        for filepath in directory.glob(f"{pattern}{ext}"):
            if filepath.is_file():
                result = import_resume(filepath, match_threshold=match_threshold)
                results.append(result)

    return results


def get_import_summary(results: list[ImportResult]) -> dict[str, Any]:
    """Summarize import results."""
    total = len(results)
    imported = sum(1 for r in results if r.resume_id and not r.skipped)
    skipped = sum(1 for r in results if r.skipped)
    errors = sum(1 for r in results if r.error)
    total_linked = sum(r.entries_linked for r in results)
    total_unmatched = sum(r.entries_unmatched for r in results)

    all_unmatched = []
    for r in results:
        all_unmatched.extend(r.unmatched_bullets)

    return {
        "total_files": total,
        "imported": imported,
        "skipped": skipped,
        "errors": errors,
        "total_entries_linked": total_linked,
        "total_entries_unmatched": total_unmatched,
        "unmatched_bullets": all_unmatched,
    }
